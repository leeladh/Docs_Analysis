import fitz
import spacy
from docx import Document
from docx.shared import Pt
import json

def extract_text_and_formatting(pdf_path):
    doc = fitz.open(pdf_path)
    extracted_data = []

    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block['type'] == 0:
                for line in block["lines"]:
                    for span in line["spans"]:
                        extracted_data.append({
                            "text": span["text"],
                            "font": span["font"],
                            "size": span["size"],
                        })

    doc.close()
    return extracted_data



def perform_ner(text_data):
    nlp = spacy.load("en_core_web_sm")
    ner_results = []

    for item in text_data:
        text = item['text']
        doc = nlp(text)
        entities = [{'entity': ent.text, 'label': ent.label_} for ent in doc.ents]
        item['entities'] = entities
        ner_results.append(item)

    return ner_results


def create_docx_with_ner(ner_results, output_path):
    # Create a new Document
    doc = Document()

    # Iterate over each item in the NER results
    for item in ner_results:
        # Create a new paragraph for each item's text
        p = doc.add_paragraph()
        run = p.add_run(item['text'])

        # Set the font size
        run.font.size = Pt(item['size'])

        # Check if there are NER entities and append them
        if item['entities']:
            entities_text = "; ".join([f"{entity['entity']} ({entity['label']})" for entity in item['entities']])
            ner_run = p.add_run("\n[Entities: " + entities_text + "]")

            # You can style the NER text differently if needed
            ner_run.font.size = Pt(item['size'] * 0.8)  # Smaller font size for entities
            ner_run.italic = True

    # Save the document to the specified output path
    doc.save(output_path)

